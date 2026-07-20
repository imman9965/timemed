#!/usr/bin/env python3
"""Convert TimesMed Markdown docs to Word (.docx) documents."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"

BLUE = RGBColor(6, 115, 222)
DARK_BLUE = RGBColor(11, 79, 138)
BODY = RGBColor(30, 30, 30)
MUTED = RGBColor(80, 80, 80)


def sanitize(text: str) -> str:
    replacements = {
        "\u2014": "-",
        "\u2013": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2026": "...",
        "\u00a0": " ",
        "\u2022": "-",
        "\u2192": "->",
        "\u2190": "<-",
        "\u2194": "<->",
        "\u2713": "[x]",
        "\u2610": "[ ]",
        "\u2611": "[x]",
        "\u20b9": "INR ",
        "\u200b": "",
        "\ufeff": "",
        "∈": "in",
        "\u2208": "in",
        "←": "<-",
        "↔": "<->",
        "≥": ">=",
        "≤": "<=",
        "×": "x",
        "•": "-",
        "–": "-",
        "—": "-",
        "₹": "INR ",
        "🔹": "",
        "🔥": "",
        "🔵": "",
        "✅": "[OK]",
        "❌": "[X]",
        "⚠️": "[WARN]",
        "📌": "",
        "💰": "",
        "📋": "",
        "🩺": "",
        "📱": "",
        "🔐": "",
        "📊": "",
        "🧾": "",
        "🏥": "",
        "💊": "",
        "🧪": "",
        "📞": "",
        "🎥": "",
        "⭐": "*",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    return text


def clean_inline(text: str) -> str:
    text = sanitize(text)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text


def set_run_font(run, size=11, bold=False, color=BODY, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)


def setup_doc(title: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = header.add_run(title)
    set_run_font(hr, size=9, bold=True, color=BLUE)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("TimesMed Health Care  |  Page ")
    set_run_font(fr, size=9, color=MUTED)
    add_page_number(footer)
    return doc


def add_cover(doc: Document, title: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_run_font(r, size=22, bold=True, color=BLUE)

    for line in (
        "TimesMed Health Care",
        "Generated from project documentation",
        "Date: 20 July 2026  |  Doc version 3.0",
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        set_run_font(r, size=11, color=MUTED)

    doc.add_paragraph()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in row):
            i += 1
            continue
        rows.append([clean_inline(c) for c in row])
        i += 1
    return rows, i


def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r_i, row in enumerate(rows):
        for c_i in range(cols):
            cell = table.rows[r_i].cells[c_i]
            text = row[c_i] if c_i < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            set_run_font(run, size=9, bold=(r_i == 0), color=DARK_BLUE if r_i == 0 else BODY)
            if r_i == 0:
                shade_cell(cell, "E8F1FC")
            elif r_i % 2 == 0:
                shade_cell(cell, "F7F9FC")
    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str], lang: str):
    label = (lang or "code").strip()
    if label.lower() == "mermaid":
        p = doc.add_paragraph()
        r = p.add_run("[Mermaid diagram — see source Markdown for editable flow]")
        set_run_font(r, size=10, bold=True, color=DARK_BLUE)

    block = sanitize("\n".join(lines))
    for line in block.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.15)
        r = p.add_run(line if line else " ")
        set_run_font(r, size=8.5, name="Consolas", color=BODY)
    doc.add_paragraph()


def render_markdown(doc: Document, md_text: str):
    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_lang = ""
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i].rstrip("\n")

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lang = line.strip()[3:]
                code_buf = []
            else:
                in_code = False
                add_code_block(doc, code_buf, code_lang)
                code_buf = []
                code_lang = ""
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if line.strip().startswith("|") and i + 1 < len(lines) and re.search(r"\|\s*-{3,}", lines[i + 1]):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        s = line.strip()
        if not s:
            i += 1
            continue

        if s.startswith("# "):
            p = doc.add_heading(clean_inline(s[2:]), level=1)
            for run in p.runs:
                set_run_font(run, size=16, bold=True, color=BLUE)
        elif s.startswith("## "):
            p = doc.add_heading(clean_inline(s[3:]), level=2)
            for run in p.runs:
                set_run_font(run, size=13, bold=True, color=DARK_BLUE)
        elif s.startswith("### "):
            p = doc.add_heading(clean_inline(s[4:]), level=3)
            for run in p.runs:
                set_run_font(run, size=12, bold=True, color=DARK_BLUE)
        elif s.startswith("#### "):
            p = doc.add_heading(clean_inline(s[5:]), level=4)
            for run in p.runs:
                set_run_font(run, size=11, bold=True, color=BODY)
        elif s.startswith("---"):
            doc.add_paragraph()
        elif s.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            r = p.add_run(clean_inline(s[2:]))
            set_run_font(r, size=10, color=DARK_BLUE)
        elif s.startswith("- [ ] ") or s.startswith("- [x] ") or s.startswith("- [X] "):
            mark = "[ ]" if "[ ]" in s[:6] else "[x]"
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(f"{mark} {clean_inline(s[6:])}")
            set_run_font(r, size=10)
        elif s.startswith("- ") or s.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(clean_inline(s[2:]))
            set_run_font(r, size=10)
        elif re.match(r"^\d+\.\s+", s):
            p = doc.add_paragraph(style="List Number")
            r = p.add_run(clean_inline(re.sub(r"^\d+\.\s+", "", s)))
            set_run_font(r, size=10)
        else:
            p = doc.add_paragraph()
            r = p.add_run(clean_inline(s))
            set_run_font(r, size=10)
        i += 1


def md_to_docx(md_path: Path, docx_path: Path, title: str):
    md_text = md_path.read_text(encoding="utf-8")
    doc = setup_doc(title)
    add_cover(doc, title)
    render_markdown(doc, md_text)
    doc.save(str(docx_path))
    print(f"Wrote {docx_path} ({docx_path.stat().st_size // 1024} KB)")


def main():
    jobs = [
        (
            DOCS / "APP_FLOW.md",
            DOCS / "TimesMed_Application_Flow.docx",
            "TimesMed — Application Flow Documentation",
        ),
        (
            DOCS / "BACKEND_REQUIREMENTS_SPECIFICATION.md",
            DOCS / "TimesMed_Backend_Requirements_Specification.docx",
            "TimesMed — Backend Requirement Specification (BRS)",
        ),
    ]
    for md, out, title in jobs:
        if not md.exists():
            print(f"Missing: {md}")
            continue
        md_to_docx(md, out, title)


if __name__ == "__main__":
    main()
